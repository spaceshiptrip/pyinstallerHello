from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Document
doc = Document()

# Add a header
header = doc.sections[0].header
header_paragraph = header.paragraphs[0]
header_paragraph.text = "This is the header text"

# Add a footer
footer = doc.sections[0].footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = "This is the footer text"

# Add a title to the document
doc.add_heading('Hello World Document', 0)

# Add a paragraph
doc.add_paragraph('This is a sample paragraph in the document.')

# Save the document
doc.save('hello_world_with_header_footer.docx')
